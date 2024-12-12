import os
from typing import Any, Dict

from langchain.schema import Document
from docx import Document as DocxDocument
from datetime import datetime
from langchain_community.tools.tavily_search import TavilySearchResults

from x_section.graph.state import GraphState
from x_section.graph.chains.write_draft_chain import write_draft_chain

web_search_tool = TavilySearchResults(k=1)


def write_draft(state: GraphState) -> Dict[str, Any]:
    print("---WRITE DRAFT (X_SECTION)---")
    summarization = state["summarization"]
    project_name = state["project_name"]
    min_length = state["min_length"]
    input_texts = state["input_texts"]
    web_search_queries = state["web_search_queries"]

    sections = []
    for idx, input_text in enumerate(input_texts):
        section_documents = [summarization]
        query = " ".join([project_name, web_search_queries[idx]])
        docs = web_search_tool.invoke({"query": query})
        web_results = "\n".join([d["content"] for d in docs])
        web_results = Document(page_content=web_results)
        if section_documents is not None:
            section_documents.append(web_results)
        else:
            section_documents = [web_results]

        response = write_draft_chain.invoke(
            {
                "documents": section_documents,
                "project_name": project_name,
                "chapter_title": input_text,
                "min_length": min_length,
            }
        )
        sections.append(response.content)

    # 创建一个新的Word文档
    doc = DocxDocument()

    # 添加章节内容
    for idx, section in enumerate(sections):
        # 添加章节标题
        doc.add_heading(f"{input_texts[idx]}", level=1)
        # 添加章节详细内容
        doc.add_paragraph(section)

    # 获取当前时间戳
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 确保文件夹路径存在
    output_dir = "files"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 保存文档
    file_path = os.path.join(output_dir, f"tender_section_{timestamp}.docx")
    doc.save(file_path)

    return {"section_path": file_path}
