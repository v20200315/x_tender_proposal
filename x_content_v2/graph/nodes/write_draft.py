import os
import json
from datetime import datetime
from typing import Any, Dict
from langchain_community.tools import TavilySearchResults
from langchain.schema import Document
from docx import Document as DocxDocument

from logger_config import logger
from x_content_v2.graph.chains.write_draft_chain import write_draft_chain
from x_content_v2.graph.state import GraphState

web_search_tool = TavilySearchResults(k=3)


def write_draft(state: GraphState) -> Dict[str, Any]:
    logger.info("---WRITE DRAFT (X_CONTENT_V2)---")
    project_name = state["project_name"]
    outline = state["outline"]
    todos = state["todos"]
    summarizations = state["summarizations"]
    images = state["images"]
    min_length = state["min_length"]

    sections = []
    for todo in todos:
        query = " ".join([project_name, todo["title"], todo["abstract"]])
        docs = web_search_tool.invoke({"query": query})
        web_results = "\n".join([d["content"] for d in docs])
        web_results = Document(page_content=web_results)
        documents = [web_results]

        logger.info(
            {
                "outline": outline,
                "title": todo["title"],
                "abstract": todo["abstract"],
                "summarizations": summarizations,
                "documents": documents,
                "images": images,
                "min_length": min_length,
            }
        )

        response = write_draft_chain.invoke(
            {
                "outline": outline,
                "title": todo["title"],
                "abstract": todo["abstract"],
                "summarizations": summarizations,
                "documents": documents,
                "images": images,
                "min_length": min_length,
            }
        )
        sections.append(response.content)

    # 创建一个新的Word文档
    doc = DocxDocument()

    # 添加章节内容
    for idx, section in enumerate(sections):
        todo = todos[idx]
        # 添加章节标题
        doc.add_heading(f"{todo["title"]}", level=1)
        # 添加章节详细内容
        doc.add_paragraph(section)

    # 获取当前时间戳
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 确保文件夹路径存在
    output_dir = "files"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 保存文档
    file_path = os.path.join(output_dir, f"tender_section_{timestamp}.docx")
    doc.save(file_path)

    # logger.info(f"summarizations: {summarizations}")
    # logger.info(f"outline: {outline}")
    # logger.info(f"todos: {todos}")
    # logger.info(f"images: {images}")
    # logger.info(f"min_length: {min_length}")

    return {"article_path": file_path}
