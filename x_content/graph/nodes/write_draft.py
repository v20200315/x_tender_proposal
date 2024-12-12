import os
import json

from typing import Any, Dict

from langchain.schema import Document
from docx import Document as DocxDocument
from datetime import datetime
from langchain_community.tools.tavily_search import TavilySearchResults

from x_content.graph.state import GraphState
from x_content.graph.chains.write_draft_chain import write_draft_chain

web_search_tool = TavilySearchResults(k=1)


def write_draft(state: GraphState) -> Dict[str, Any]:
    print("---WRITE DRAFT (X_CONTENT)---")
    summarization = state["summarization"]
    outline_json = json.loads(state["outline_json"])
    project_name = outline_json["project_name"]
    article_json = outline_json
    for chapter in article_json["chapters"]:
        chapter_documents = [summarization]
        query = " ".join([project_name, chapter["query"]])
        docs = web_search_tool.invoke({"query": query})
        web_results = "\n".join([d["content"] for d in docs])
        web_results = Document(page_content=web_results)
        if chapter_documents is not None:
            chapter_documents.append(web_results)
        else:
            chapter_documents = [web_results]

        article = write_draft_chain.invoke(
            {
                "documents": chapter_documents,
                "project_name": project_name,
                "chapter_title": chapter["title"],
                "chapter_content": chapter["content"],
                "min_length": 600,
            }
        )
        chapter["article"] = article.content

    # 创建一个新的Word文档
    doc = DocxDocument()

    # 添加标题
    doc.add_heading(f"{article_json['project_name']} 投标文件", 0)

    # 添加章节内容
    for chapter in article_json["chapters"]:
        # 添加章节标题
        doc.add_heading(f"第{chapter['order']}章: {chapter['title']}", level=1)
        # 添加章节详细内容
        if "article" in chapter:
            doc.add_paragraph(chapter["article"])

    # 获取当前时间戳
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 确保文件夹路径存在
    output_dir = "files"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 保存文档
    file_path = os.path.join(output_dir, f"tender_document_{timestamp}.docx")
    doc.save(file_path)

    return {"article_path": file_path}
