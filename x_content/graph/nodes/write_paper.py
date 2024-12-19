import os
from datetime import datetime
from typing import Any, Dict
from docx import Document as DocxDocument

from logger_config import logger
from x_content.graph.state import GraphState


def write_paper(state: GraphState) -> Dict[str, Any]:
    logger.info("---WRITE PAPER (X_CONTENT)---")
    done_list = state["done_list"]

    # 创建一个新的Word文档
    doc = DocxDocument()

    for done in done_list:
        # 添加章节标题
        doc.add_heading(f"{done["title"]}", level=1)
        # 添加章节详细内容
        doc.add_paragraph(done["content"])

    # 获取当前时间戳
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 确保文件夹路径存在
    output_dir = "files"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 保存文档
    file_path = os.path.join(output_dir, f"tender_section_{timestamp}.docx")
    doc.save(file_path)
    return {"article_path": file_path}
